"""In-memory PDF and DOCX builders for CV extraction tests."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from pypdf import PdfReader, PdfWriter


def build_text_pdf(pages: list[list[str]]) -> bytes:
    """Build a minimal text-readable PDF with one content stream per page."""
    if not pages:
        pages = [[]]
    objects: list[str] = []
    objects.append("<< /Type /Catalog /Pages 2 0 R >>")
    page_ids = [4 + 2 * index for index in range(len(pages))]
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>")
    objects.append("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for index, lines in enumerate(pages):
        stream = _page_stream(lines)
        objects.append(
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Contents {5 + 2 * index} 0 R /Resources << /Font << /F1 3 0 R >> >> >>"
        )
        content_length = len(stream.encode("latin-1"))
        objects.append(f"<< /Length {content_length} >>\nstream\n{stream}\nendstream")
    return _assemble_pdf(objects)


def build_blank_pdf(page_count: int) -> bytes:
    """Build a PDF with empty pages and no extractable text."""
    return build_text_pdf([[] for _ in range(page_count)])


def encrypt_pdf(file_bytes: bytes, password: str = "secret") -> bytes:
    """Return an encrypted copy of a PDF."""
    writer = PdfWriter()
    writer.append(PdfReader(BytesIO(file_bytes)))
    writer.encrypt(password)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def build_ordered_docx() -> bytes:
    """Build a DOCX with a paragraph, table, and trailing paragraph."""
    document = Document()
    document.add_paragraph("Intro paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "cell a"
    table.cell(0, 1).text = "cell b"
    document.add_paragraph("Closing paragraph")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_text_docx(
    *,
    paragraphs: list[str] | None = None,
    table: list[list[str]] | None = None,
) -> bytes:
    """Build a DOCX with paragraphs and an optional table in document order."""
    document = Document()
    for text in paragraphs or []:
        document.add_paragraph(text)
    if table:
        rows = len(table)
        cols = max(len(row) for row in table)
        created = document.add_table(rows=rows, cols=cols)
        for row_index, row in enumerate(table):
            for cell_index, value in enumerate(row):
                created.cell(row_index, cell_index).text = value
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _page_stream(lines: list[str]) -> str:
    commands = ["BT /F1 12 Tf 72 720 Td"]
    for index, line in enumerate(lines):
        if index:
            commands.append("0 -16 Td")
        commands.append(f"({_pdf_escape(line)}) Tj")
    commands.append("ET")
    return " ".join(commands)


def _pdf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _assemble_pdf(object_bodies: list[str]) -> bytes:
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(object_bodies, start=1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n{body}\nendobj\n".encode("latin-1"))
    xref_at = len(output)
    size = len(object_bodies) + 1
    output.extend(f"xref\n0 {size}\n".encode("latin-1"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
    output.extend(
        f"trailer\n<< /Size {size} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n".encode("latin-1")
    )
    return bytes(output)
